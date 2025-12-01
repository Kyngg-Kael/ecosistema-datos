import io
import pandas as pd
import geopandas as gpd
import matplotlib.pyplot as plt
import contextily as ctx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from shapely.geometry import Polygon, MultiPolygon

def create_chart_image(plot_func, *args, **kwargs):
    """Helper para convertir gráficos a bytes para Word."""
    buf = io.BytesIO()
    fig = plt.figure(figsize=(6, 4), dpi=150)
    plot_func(fig, *args, **kwargs)
    plt.tight_layout()
    plt.savefig(buf, format='png', bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf

# ===================== FUNCIÓN DE MAPA (CON ZOOM ALEJADO) =====================
def plot_polygon_outline(fig, geometry):
    ax = fig.add_subplot(111)
    
    # 1. Crear GeoSeries y reproyectar a Web Mercator
    gs = gpd.GeoSeries([geometry], crs="EPSG:4326")
    gs_3857 = gs.to_crs(epsg=3857)
    
    # 2. Dibujar el polígono (Azul translúcido)
    gs_3857.plot(ax=ax, color='#2E86AB', alpha=0.4, edgecolor='#2E86AB', linewidth=2)
    
    # 3. CALCULAR ZOOM (ALEJAR LA CÁMARA)
    minx, miny, maxx, maxy = gs_3857.total_bounds
    
    # Calculamos el ancho y alto del polígono
    width = maxx - minx
    height = maxy - miny
    
    margin_factor = 0.8 
    
    ax.set_xlim(minx - (width * margin_factor), maxx + (width * margin_factor))
    ax.set_ylim(miny - (height * margin_factor), maxy + (height * margin_factor))
    
    # 4. AÑADIR MAPA BASE
    try:
        ctx.add_basemap(ax, source=ctx.providers.CartoDB.Positron)
    except Exception:
        pass

    ax.axis('off')
    ax.set_title("Contexto Geográfico Regional", fontsize=10)

# ===================== OTRAS GRÁFICAS =====================
def plot_forest_pie(fig, df):
    ax = fig.add_subplot(111)
    if not df.empty:
        df_plot = df.copy()
        colors = ['#4CAF50', '#FFC107', '#F44336', '#2196F3', '#9C27B0']
        wedges, texts, autotexts = ax.pie(
            df_plot['Área (ha)'], 
            labels=df_plot['Leyenda'], 
            autopct='%1.1f%%', 
            startangle=90, 
            colors=colors[:len(df_plot)]
        )
        plt.setp(autotexts, size=8, weight="bold", color="white")
        ax.set_title("Distribución de Coberturas (IDEAM)", fontsize=10)

def plot_biodiversity_bar(fig, df):
    ax = fig.add_subplot(111)
    if not df.empty:
        df_sorted = df.sort_values(by='Especies (GBIF)', ascending=True)
        bars = ax.barh(df_sorted['Grupo'], df_sorted['Especies (GBIF)'], color='#2E86AB')
        ax.bar_label(bars, padding=3)
        ax.set_xlabel("N° Especies Registradas")
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.set_title("Riqueza Potencial por Grupo (GBIF)", fontsize=10)

# ===================== GENERADOR DE REPORTE =====================
def generate_docx_report(context):
    doc = Document()
    
    # Datos
    geometry = context.get('geometry')
    loc = context.get('location_info', {})
    raster = context.get('raster_data')
    vectors = context.get('vector_data', {})
    bio = context.get('biodiversity_data')
    sat = context.get('satellite_data', {})

    # --- PORTADA ---
    title = doc.add_heading('Bitácora Territorial del Siglo XXI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Tierra, Vida y Paz')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph("\n")

    # --- CROQUIS (MAPA) ---
    if geometry:
        doc.add_heading('Ubicación del Territorio', level=1)
        
        # Generamos el mapa
        img_stream = create_chart_image(plot_polygon_outline, geometry)
        
        doc.add_picture(img_stream, width=Inches(4.0))
        
        # Centramos la imagen
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Info de ubicación debajo del mapa
        muni = loc.get('municipio', 'No determinado')
        depto = loc.get('departamento', '')
        
        p_info = doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_info.add_run(f"\nMunicipio: ").bold = True
        p_info.add_run(f"{muni}, {depto}  |  ")
        p_info.add_run(f"Fecha: ").bold = True
        p_info.add_run(f"{pd.Timestamp.now().strftime('%Y-%m-%d')}")

    # --- RESUMEN EJECUTIVO ---
    doc.add_heading('Resumen Ejecutivo', level=1)
    
    bosque_ha = raster[raster['Leyenda'].str.contains("Bosque", case=False)]['Área (ha)'].sum() if raster is not None and not raster.empty else 0
    co2_total = sat['biomass']['stats'].get('Captura Potencial CO2 (Mg)', 0) if sat.get('biomass') else 0
    spp_total = bio['Especies (GBIF)'].sum() if bio is not None else 0

    resumen = (
        f"Este reporte presenta la caracterización integral del predio ubicado en {muni}. "
        f"El análisis Multidimensional de IA identifica un área con una cobertura boscosa de {bosque_ha:.1f} hectáreas, "
        f"albergando un potencial de biodiversidad de {spp_total} especies registradas históricamente.\n\n"
        f"Desde la perspectiva climática, se estima un potencial de captura de carbono de {co2_total:,.0f} toneladas de CO2e. "
        "El cruce con bases de datos oficiales (SIPRA) permite identificar el contexto legal y productivo del territorio."
    )
    doc.add_paragraph(resumen)

    # --- CAP 1: AMBIENTAL ---
    doc.add_heading('1. Componente Ambiental (IDEAM)', level=1)
    if raster is not None and not raster.empty:
        col1, col2 = doc.add_table(rows=1, cols=2).rows[0].cells
        img_stream = create_chart_image(plot_forest_pie, raster)
        run = col1.paragraphs[0].add_run()
        run.add_picture(img_stream, width=Inches(3.0))
        
        t_data = col2.add_table(rows=1, cols=2)
        t_data.style = 'Table Grid'
        hdr = t_data.rows[0].cells
        hdr[0].text = 'Cobertura'; hdr[1].text = 'Área (ha)'
        for _, row in raster.iterrows():
            r = t_data.add_row().cells
            r[0].text = str(row['Leyenda']); r[1].text = f"{row['Área (ha)']:.1f}"
    else:
        doc.add_paragraph("Sin datos de cobertura boscosa.")

    # --- CAP 2: LEGAL ---
    doc.add_heading('2. Contexto Legal y Productivo (SIPRA)', level=1)
    if vectors:
        for titulo, df in vectors.items():
            doc.add_heading(titulo, level=2)
            cats = df['Categoría'].tolist()
            doc.add_paragraph(f"Categorías identificadas: {', '.join(cats[:5])}.")
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light List Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Categoría'; hdr[1].text = 'Área (ha)'
            for c, a in zip(cats, df['area_total_ha']):
                row = table.add_row().cells
                row[0].text = str(c); row[1].text = f"{a:.1f}"
    else:
        doc.add_paragraph("El área no presenta intersecciones con las capas legales analizadas.")

    # --- CAP 3: SATELITAL ---
    doc.add_page_break()
    doc.add_heading('3. Inteligencia Satelital (IA)', level=1)
    
    if sat.get('biomass'):
        s = sat['biomass']['stats']
        doc.add_heading('Biomasa y Mercado de Carbono (GEDI)', level=2)
        p = doc.add_paragraph("Estimación basada en tecnología LiDAR satelital y Machine Learning:\n")
        p.add_run(f"• Biomasa Media: ").bold = True; p.add_run(f"{s.get('Media (Mg/ha)')} Mg/ha\n")
        p.add_run(f"• Potencial CO2e: ").bold = True; p.add_run(f"{s.get('Captura Potencial CO2 (Mg)')} Toneladas").font.color.rgb = RGBColor(46, 134, 171)

    if sat.get('canopy'):
        s = sat['canopy']['stats']
        doc.add_heading('Altura del Dosel (Meta AI)', level=2)
        doc.add_paragraph(f"La altura promedio del dosel vegetal en el área es de {s.get('Promedio (m)')} metros.")

    # --- CAP 4: BIODIVERSIDAD ---
    doc.add_heading('4. Riqueza de Especies (GBIF)', level=1)
    if bio is not None and not bio.empty:
        doc.add_paragraph(f"Registros históricos en el área: {spp_total} especies.")
        img_stream = create_chart_image(plot_biodiversity_bar, bio)
        doc.add_picture(img_stream, width=Inches(5.0))
    else:
        doc.add_paragraph("No se encontraron registros biológicos directos.")

    # --- CIERRE EQUIPO ---
    doc.add_paragraph("\n\n")
    p_final = doc.add_paragraph("Reporte generado por:")
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_team = doc.add_paragraph("Equipo Datos al Ecosistema: Carlos Betancur, Paula Castro, Mario Ortegon y Santiago Restrepo")
    p_team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_team.runs[0].bold = True
    p_team.style = 'Quote'

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output